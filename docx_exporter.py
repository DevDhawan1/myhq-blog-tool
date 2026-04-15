"""
Converts a generated blog result dict into a formatted .docx file.
"""
from io import BytesIO
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_heading_style(para, level: int):
    run = para.runs[0] if para.runs else para.add_run()
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0f, 0x3d, 0x60)


def _add_meta_row(table, label: str, value: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[1].text = value


def _parse_html_to_doc(doc: Document, html: str):
    """
    Simple HTML → docx parser. Handles h2, h3, p, ul/li, a tags.
    Strips remaining tags and writes plain paragraphs.
    """
    # Split on block-level tags (including tables and styled divs)
    blocks = re.split(r'(<h[23][^>]*>.*?</h[23]>|<p[^>]*>.*?</p>|<ul[^>]*>.*?</ul>|<table[^>]*>.*?</table>|<div[^>]*>.*?</div>)', html, flags=re.DOTALL)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # H2
        m = re.match(r'<h2[^>]*>(.*?)</h2>', block, re.DOTALL)
        if m:
            text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
            if text:
                para = doc.add_paragraph(text, style="Heading 2")
            continue

        # H3
        m = re.match(r'<h3[^>]*>(.*?)</h3>', block, re.DOTALL)
        if m:
            text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
            if text:
                para = doc.add_paragraph(text, style="Heading 3")
            continue

        # UL list
        m = re.match(r'<ul[^>]*>(.*?)</ul>', block, re.DOTALL)
        if m:
            items = re.findall(r'<li[^>]*>(.*?)</li>', m.group(1), re.DOTALL)
            for item in items:
                text = re.sub(r'<[^>]+>', '', item).strip()
                if text:
                    doc.add_paragraph(text, style="List Bullet")
            continue

        # Paragraph (with possible <a> links)
        m = re.match(r'<p[^>]*>(.*?)</p>', block, re.DOTALL)
        if m:
            inner = m.group(1)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Split by <a> tags to preserve link text
            parts = re.split(r'(<a\s[^>]*>.*?</a>)', inner, flags=re.DOTALL)
            for part in parts:
                link_m = re.match(r'<a\s[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', part, re.DOTALL)
                if link_m:
                    link_text = re.sub(r'<[^>]+>', '', link_m.group(2)).strip()
                    if link_text:
                        run = para.add_run(link_text)
                        run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
                        run.font.underline = True
                else:
                    text = re.sub(r'<[^>]+>', '', part)
                    if text.strip():
                        para.add_run(text)
            continue

        # Definition box (styled div)
        m = re.match(r'<div[^>]*class="definition-box"[^>]*>(.*?)</div>', block, re.DOTALL)
        if m:
            text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
            if text:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.3)
                run = para.add_run(text)
                run.font.size = Pt(11)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            continue

        # Quick-recap / other styled divs
        m = re.match(r'<div[^>]*>(.*?)</div>', block, re.DOTALL)
        if m:
            text = re.sub(r'<[^>]+>', '', m.group(1)).strip()
            if text:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.3)
                run = para.add_run(text)
                run.font.size = Pt(10)
                run.font.italic = True
            continue

        # Table
        m = re.match(r'<table[^>]*>(.*?)</table>', block, re.DOTALL)
        if m:
            rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', m.group(1), re.DOTALL)
            parsed_rows = []
            max_cols = 0
            for row_html in rows_html:
                cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row_html, re.DOTALL)
                cell_texts = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
                parsed_rows.append(cell_texts)
                max_cols = max(max_cols, len(cell_texts))
            if parsed_rows and max_cols > 0:
                tbl = doc.add_table(rows=len(parsed_rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row_data in enumerate(parsed_rows):
                    for c_idx in range(max_cols):
                        cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                        tbl.cell(r_idx, c_idx).text = cell_text
                        if r_idx == 0:  # bold header row
                            for run in tbl.cell(r_idx, c_idx).paragraphs[0].runs:
                                run.font.bold = True
                doc.add_paragraph()
            continue

        # Fallback: strip all tags and add as plain paragraph
        text = re.sub(r'<[^>]+>', '', block).strip()
        if text:
            doc.add_paragraph(text)


# ── Main export function ─────────────────────────────────────────────────────

def build_docx(result: dict, image_bytes: bytes | None = None) -> bytes:
    """
    Takes the generated blog result dict and returns a .docx as bytes.
    """
    doc = Document()

    # ── Page title ────────────────────────────────────────────────────────────
    title_para = doc.add_heading(result.get("blog_title", "Blog Draft"), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Feature image ─────────────────────────────────────────────────────────
    if image_bytes:
        from io import BytesIO as _BytesIO
        doc.add_paragraph()
        img_stream = _BytesIO(image_bytes)
        doc.add_picture(img_stream, width=Inches(6))
        cap = doc.add_paragraph("Feature Image")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── SEO metadata table ────────────────────────────────────────────────────
    doc.add_heading("SEO Metadata", level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)

    _add_meta_row(table, "Meta Title", result.get("meta_title", ""))
    _add_meta_row(table, "Meta Description", result.get("meta_description", ""))
    _add_meta_row(table, "Focus Keyword", result.get("focus_keyword", ""))
    _add_meta_row(table, "Subsidiary Keywords", ", ".join(result.get("subsidiary_keywords", [])))
    _add_meta_row(table, "URL Slug", result.get("url_slug", ""))

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Key Takeaways (TL;DR) ─────────────────────────────────────────────────
    tl_dr = result.get("tl_dr", [])
    if tl_dr:
        doc.add_heading("Key Takeaways", level=2)
        for item in tl_dr:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph()

    # ── Blog content ──────────────────────────────────────────────────────────
    doc.add_heading("Blog Content", level=2)
    _parse_html_to_doc(doc, result.get("content", ""))

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
